import contextlib
import json
import os
import shutil
import sys
import time

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")


# 빈칸 번호용: ▁1▁ 형식 (문단별 번호 restart)
def blank_format(num):
    """빈칸 번호를 ▁1▁ 형식으로 반환"""
    return f"▁{num}▁"


def load_law_terms():
    db_path = os.path.join(os.path.dirname(__file__), "law_terms_db.json")
    with open(db_path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_blank_words_for_file(filename):
    db = load_law_terms()

    if "국가배상" in filename:
        words = []
        words.extend(db.get("국가배상법", {}).get("핵심법조문", []))
        words.extend(db.get("국가배상법", {}).get("공무원", []))
        words.extend(db.get("국가배상법", {}).get("요건", []))
        words.extend(db.get("국가배상법", {}).get("직무범위", []))
        words.extend(db.get("국가배상법", {}).get("외관이론", []))
        words.extend(db.get("국가배상법", {}).get("과실", []))
        words.extend(db.get("국가배상법", {}).get("법령위반", []))
        words.extend(db.get("국가배상법", {}).get("공공단체", []))
        words.extend(db.get("국가배상법", {}).get("부작위", []))
        words.extend(db.get("국가배상법", {}).get("판결", []))
        words.extend(db.get("국가배상법", {}).get("영조물", []))
        words.extend(db.get("국가배상법", {}).get("이중배상금지", []))
        words.extend(db.get("국가배상법", {}).get("소멸시효", []))
        words.extend(db.get("국가배상법", {}).get("비용부담자", []))
        return list(set(words))

    elif "지자체" in filename or "공무원법" in filename:
        words = []
        words.extend(db.get("지방자치단체_공무원법", {}).get("지방자치", []))
        words.extend(db.get("지방자치단체_공무원법", {}).get("공무원", []))
        words.extend(db.get("지방자치단체_공무원법", {}).get("인사", []))
        words.extend(db.get("지방자치단체_공무원법", {}).get("직무", []))
        words.extend(db.get("토지보상법", {}).get("총칙", []))
        words.extend(db.get("토지보상법", {}).get("보상", []))
        words.extend(db.get("토지보상법", {}).get("절차", []))
        words.extend(db.get("토지보상법", {}).get("특례", []))
        words.extend(db.get("재개발_도시정비", {}).get("총칙", []))
        words.extend(db.get("재개발_도시정비", {}).get("절차", []))
        words.extend(db.get("재개발_도시정비", {}).get("동의", []))
        words.extend(db.get("재개발_도시정비", {}).get("인가", []))
        return list(set(words))

    elif "행정과 법원리" in filename or "행정절차" in filename:
        words = []
        words.extend(db.get("행정법_일반원칙", {}).get("비례원칙", []))
        words.extend(db.get("행정법_일반원칙", {}).get("평등원칙", []))
        words.extend(db.get("행정법_일반원칙", {}).get("신뢰보호", []))
        words.extend(db.get("행정법_일반원칙", {}).get("기타원칙", []))
        words.extend(db.get("행정절차법", {}).get("총칙", []))
        words.extend(db.get("행정절차법", {}).get("원칙", []))
        words.extend(db.get("행정절차법", {}).get("절차", []))
        words.extend(db.get("행정절차법", {}).get("행정지도", []))
        return list(set(words))

    elif "행정심판" in filename or "행정소송" in filename:
        words = []
        words.extend(db.get("행정심판법", {}).get("이의신청", []))
        words.extend(db.get("행정심판법", {}).get("행정심판", []))
        words.extend(db.get("행정심판법", {}).get("취소심판", []))
        words.extend(db.get("행정심판법", {}).get("형성재결", []))
        words.extend(db.get("행정심판법", {}).get("재결효력", []))
        words.extend(db.get("행정심판법", {}).get("절차", []))
        words.extend(db.get("행정소송법", {}).get("소송유형", []))
        words.extend(db.get("행정소송법", {}).get("청구요건", []))
        words.extend(db.get("행정소송법", {}).get("기판력", []))
        words.extend(db.get("행정소송법", {}).get("기속력", []))
        words.extend(db.get("행정소송법", {}).get("집행정지", []))
        words.extend(db.get("행정소송법", {}).get("절차", []))
        return list(set(words))

    elif "행정입법" in filename or "행정계획" in filename:
        words = []
        words.extend(db.get("행정입법", {}).get("법원리", []))
        words.extend(db.get("행정입법", {}).get("구분", []))
        words.extend(db.get("행정입법", {}).get("법령보충규칙", []))
        words.extend(db.get("행정입법", {}).get("행정계획", []))
        return list(set(words))

    return []


def replace_word_in_para(para, word, new_text):
    """문단에서 단어를 찾아서 교체 (run 분할 처리)"""
    full_text = para.text
    if word not in full_text:
        return False

    # 전체 텍스트에서 위치 찾기
    start_idx = full_text.find(word)
    if start_idx == -1:
        return False

    # run별로 텍스트 위치 추적
    current_pos = 0
    for run in para.runs:
        run_start = current_pos
        run_end = current_pos + len(run.text)

        # 이 run에 교체할 단어가 포함된 경우
        if run_start <= start_idx < run_end:
            # run 텍스트에서 단어 위치
            word_start_in_run = start_idx - run_start
            word_end_in_run = word_start_in_run + len(word)

            # run 분할: 앞부분 + 빈칸 + 뒷부분
            before = run.text[:word_start_in_run]
            after = run.text[word_end_in_run:]
            run.text = before + new_text + after
            return True

        current_pos = run_end

    return False


def get_heading_text(para):
    """Heading 스타일에서 제목 텍스트 추출"""
    if (
        para.style
        and para.style.name
        and ("Heading" in para.style.name or "제목" in para.style.name)
    ):
        return para.text.strip()
    return None


def create_quiz_from_db(source_path, output_dir):
    """
    DB에서 빈칸 단어를 가져와서 문제 생성
    - 문단별 번호 restart (1번부터)
    - 답안지에 문단 제목 + 번호
    """
    filename = os.path.splitext(os.path.basename(source_path))[0]
    blank_words = get_blank_words_for_file(filename)

    if not blank_words:
        print(f"{filename}: 해당하는 빈칸 단어 없음")
        return

    # 1. 원본 파일 복사
    question_path = os.path.join(output_dir, f"{filename}_문제.docx")

    if os.path.exists(question_path):
        with contextlib.suppress(PermissionError):
            os.remove(question_path)

    for _attempt in range(3):
        try:
            shutil.copy2(source_path, question_path)
            break
        except PermissionError:
            time.sleep(1)
    else:
        print(f"{filename}: 파일 복사 실패")
        return

    # 2. 복사된 파일 열기
    doc = docx.Document(question_path)

    # 3. 섹션 추적 및 답안용 데이터
    current_section = "기타"
    answer_list = []  # [(문단번호, 빈칸번호, 단어, 섹션제목), ...]
    section_blanks = {}  # {섹션제목: [(문단번호, 빈칸번호, 단어), ...]}
    para_num = 0  # 전체 문단 번호

    # 4. 문단별 빈칸 변경
    for para in doc.paragraphs:
        para_num += 1

        # 섹션 변경 감지
        heading = get_heading_text(para)
        if heading:
            current_section = heading

        # 빈칸 변경 (문단당 최대 2개, 번호 restart)
        blank_count = 0
        for word in blank_words:
            if blank_count >= 2:
                break
            if "▁" in para.text:
                break
            if word in para.text:
                blank_count += 1
                fmt = blank_format(blank_count)
                if replace_word_in_para(para, word, fmt):
                    answer_list.append((para_num, blank_count, word, current_section))
                    if current_section not in section_blanks:
                        section_blanks[current_section] = []
                    section_blanks[current_section].append(
                        (para_num, blank_count, word)
                    )

    # 5. 표에서도 빈칸 변경
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    heading = get_heading_text(para)
                    if heading:
                        current_section = heading

                    blank_count = 0
                    for word in blank_words:
                        if blank_count >= 2:
                            break
                        if "▁" in para.text:
                            break
                        if word in para.text:
                            blank_count += 1
                            fmt = blank_format(blank_count)
                            if replace_word_in_para(para, word, fmt):
                                answer_list.append(
                                    (para_num, blank_count, word, current_section)
                                )
                                if current_section not in section_blanks:
                                    section_blanks[current_section] = []
                                section_blanks[current_section].append(
                                    (para_num, blank_count, word)
                                )

    # 6. 저장
    doc.save(question_path)

    # 7. 답안 파일 생성 (섹션별 그룹핑)
    answer_path = os.path.join(output_dir, f"{filename}_답안.docx")
    answer_doc = docx.Document()

    title_para = answer_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{filename} 답안")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"

    for section, blanks in section_blanks.items():
        # 섹션 제목
        section_para = answer_doc.add_paragraph()
        section_run = section_para.add_run(f"【{section}】")
        section_run.bold = True
        section_run.font.size = Pt(12)
        section_run.font.name = "Arial"

        # 빈칸 목록 (문단번호-빈칸번호)
        for para_no, blank_no, word in blanks:
            item_para = answer_doc.add_paragraph()
            item_run = item_para.add_run(f"  {para_no}-{blank_no}. {word}")
            item_run.font.name = "Arial"

    answer_doc.save(answer_path)

    print(
        f"생성 완료: {os.path.basename(question_path)}, {os.path.basename(answer_path)}"
    )
    print(f"  빈칸 총 {len(answer_list)}개, 섹션 {len(section_blanks)}개")


if __name__ == "__main__":
    base_dir = r"D:\code proj\word proj"

    files = [
        "사례_국가배상, 손실보상",
        "사례_지자체, 공무원법, 토지보상법, 재개발 등",
        "사례_행정과 법원리, 행정절차",
        "사례_행정심판, 행정소송",
        "사례_행정입법, 행정계획",
    ]

    for f in files:
        source = os.path.join(base_dir, f"{f}.docx")
        if os.path.exists(source):
            create_quiz_from_db(source, base_dir)
        else:
            print(f"{f}.docx 파일 없음")
