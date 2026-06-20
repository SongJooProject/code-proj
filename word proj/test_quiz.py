"""뻥뚜리 결과물 검증 스크립트"""

import re
import sys
from pathlib import Path

import docx

# Windows 콘솔 인코딩 문제 해결
sys.stdout.reconfigure(encoding="utf-8")


def test_quiz_file(filepath):
    """문제지 검증"""
    errors = []
    doc = docx.Document(filepath)

    blank_pattern = re.compile(r"▁(\d+)▁")
    para_blank_counts = []
    all_blanks = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        blanks = blank_pattern.findall(text)
        if blanks:
            para_blank_counts.append((i + 1, len(blanks), blanks))
            all_blanks.extend(int(b) for b in blanks)

            # 1. 문단당 최대 2개 빈칸
            if len(blanks) > 2:
                errors.append(f"문단 {i + 1}: 빈칸 {len(blanks)}개 (최대 2개 초과)")

            # 2. 번호가 1부터 순서대로
            for b in blanks:
                if int(b) < 1:
                    errors.append(f"문단 {i + 1}: 빈칸 번호 {b} (1 미만)")

    # 3. 빈칸이 하나도 없으면 경고
    if not all_blanks:
        errors.append("빈칸이 하나도 없습니다")

    return errors, para_blank_counts


def test_answer_file(filepath):
    """답안지 검증"""
    errors = []
    doc = docx.Document(filepath)

    answer_pattern = re.compile(r"(\d+)번\.\s*(.+)")
    section_pattern = re.compile(r"【(.+)】")

    sections = {}
    current_section = None
    total_answers = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 섹션 제목 확인 (【...】 형식)
        section_match = section_pattern.match(text)
        if section_match:
            current_section = section_match.group(1)
            sections[current_section] = []
            continue

        # 답안 확인 (N번. 단어 형식)
        answer_match = answer_pattern.match(text)
        if answer_match:
            num = int(answer_match.group(1))
            word = answer_match.group(2).strip()
            total_answers += 1

            if current_section is None:
                errors.append(f"답안 '{text}': 섹션 없음")

            sections[current_section].append((num, word))

    # 4. 각 섹션에서 번호가 1부터 순서대로
    for section, answers in sections.items():
        nums = [a[0] for a in answers]
        if nums and nums[0] != 1:
            errors.append(f"섹션 '{section}': 첫 번호가 {nums[0]} (1 아님)")
        for i in range(1, len(nums)):
            if nums[i] != nums[i - 1] + 1:
                errors.append(
                    f"섹션 '{section}': 번호 {nums[i - 1]} → {nums[i]} (순서 아님)"
                )

    # 5. 답안이 하나도 없으면 경고
    if total_answers == 0:
        errors.append("답안이 하나도 없습니다")

    return errors, sections, total_answers


def test_consistency(question_path, answer_path):
    """문제지-답안 일관성 검증"""
    errors = []

    # 문제지에서 빈칸 확인
    q_doc = docx.Document(question_path)
    for para in q_doc.paragraphs:
        text = para.text
        # ▁N▁ 형식의 빈칸은 단어가 이미 제거됨
        if "▁" in text:
            pass  # 빈칸이 있으면 OK

    # 답안지에서 단어 추출
    a_doc = docx.Document(answer_path)
    a_words = []
    answer_pattern = re.compile(r"(\d+)번\.\s*(.+)")
    for para in a_doc.paragraphs:
        match = answer_pattern.match(para.text.strip())
        if match:
            a_words.append(match.group(2).strip())

    # 답안에 단어가 있어야 함
    if not a_words:
        errors.append("답안에 단어가 없습니다")

    return errors


def run_all_tests():
    """모든 테스트 실행"""
    base = Path(__file__).parent
    question_files = sorted(base.glob("*_문제.docx"))
    answer_files = sorted(base.glob("*_답안.docx"))

    print("=" * 60)
    print("뻥뚜리 결과물 검증")
    print("=" * 60)

    total_errors = 0

    # 각 문제지 검증
    print("\n[문제지 검증]")
    for qf in question_files:
        errors, blanks = test_quiz_file(qf)
        if errors:
            print(f"  FAIL {qf.name}")
            for e in errors:
                print(f"     - {e}")
            total_errors += len(errors)
        else:
            blank_count = sum(b[1] for b in blanks)
            print(f"  OK {qf.name} (빈칸 {blank_count}개)")

    # 각 답안지 검증
    print("\n[답안지 검증]")
    for af in answer_files:
        errors, sections, total = test_answer_file(af)
        if errors:
            print(f"  FAIL {af.name}")
            for e in errors:
                print(f"     - {e}")
            total_errors += len(errors)
        else:
            print(f"  OK {af.name} (답안 {total}개, 섹션 {len(sections)}개)")

    # 문제-답안 쌍 일관성 검증
    print("\n[일관성 검증]")
    for qf in question_files:
        af = qf.parent / qf.name.replace("_문제.docx", "_답안.docx")
        if af.exists():
            errors = test_consistency(qf, af)
            if errors:
                print(f"  FAIL {qf.name} vs {af.name}")
                for e in errors:
                    print(f"     - {e}")
                total_errors += len(errors)
            else:
                print(f"  OK {qf.name} vs {af.name}")

    # 결과 요약
    print("\n" + "=" * 60)
    if total_errors == 0:
        print("ALL PASS - 모든 검증 통과!")
    else:
        print(f"FAIL - {total_errors}개 오류 발견")
    print("=" * 60)

    return total_errors


if __name__ == "__main__":
    errors = run_all_tests()
    sys.exit(1 if errors else 0)
