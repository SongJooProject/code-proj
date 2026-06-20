import os
import shutil
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")


# 각 문단별 빈칸 단어 목록 (문단 인덱스: [빈칸 단어들])
# 문단 인덱스는 원본 paragraphs 리스트 기준
blank_words_by_paragraph = {
    0: [
        "직무상 불법행위",
        "손해",
        "국민",
        "법률",
        "국가",
        "공공단체",
        "정당한 배상",
        "청구",
    ],
    2: ["군인", "군무원", "경찰공무원", "전투", "훈련", "손해", "보상"],
    5: ["국가배상청구소송", "민사소송", "행정소송", "당사자소송", "공권"],
    8: ["공무원", "공무위탁", "사인", "공행정사무", "이행보조자"],
    12: [
        "권력작용",
        "관리작용",
        "광의설",
        "협의설",
        "사경제주체",
        "권력적",
        "비권력적",
    ],
    14: ["직무행위", "의사", "객관적", "외관", "판단"],
    16: ["직무집행", "과실", "공무원", "평균적", "주의의무"],
    18: ["법령", "위반", "행위의무", "인권존중", "권력남용금지", "신의성실", "준칙"],
    27: ["공공단체", "공무", "수탁", "불법행위", "손해배상책임", "국가"],
    28: ["공공단체", "공법인", "행정주체", "권한", "위탁", "배상책임"],
    29: [
        "공법인",
        "국가",
        "위탁",
        "공행정사무",
        "임직원",
        "피용인",
        "고의",
        "과실",
        "법령",
    ],
    31: ["헌법", "국가", "공공단체", "배상책임", "국가배상법", "민법"],
    33: ["공무원", "직무", "관계법규", "지식", "법규해석", "행정처분", "과실"],
    36: ["상대적 위법성설", "피침해이익", "침해", "가해행위", "객관적", "정당성"],
    37: ["광의의 행위위법성설", "항고소송", "위법", "국가", "손해방지의무"],
    38: [
        "행위위법성설",
        "상대적위법성설",
        "취소소송",
        "국가배상청구소송",
        "목적",
        "판단기준",
    ],
    39: ["객관적", "정당성", "피침해이익", "침해행위", "피해자", "손해정도"],
    40: ["권한", "불행사", "부작위", "위법성", "작위의무", "사익보호성"],
    42: ["작위의무", "기속행위", "재량행위", "수축", "응답의무"],
    43: ["부작위", "법령위반", "신청", "작위의무"],
    44: ["법령", "작위의무", "국가", "생명", "신체", "재산", "위험배제"],
    46: [
        "부작위",
        "위법성",
        "법익",
        "손해",
        "심각성",
        "공무원",
        "예견가능성",
        "회피가능성",
    ],
    48: [
        "행정기관",
        "작위",
        "권한행사",
        "이익",
        "반사적 이익",
        "국가배상책임",
        "사익보호성",
    ],
    50: ["사익보호성", "법령", "의무", "공공일반", "행정기관", "내부질서", "인과관계"],
    51: ["재판", "독립성", "확정판결", "기판력", "국가배상책임"],
    52: ["법관", "위법", "부당", "목적", "재판", "기준", "권한"],
    53: ["재판", "불복절차", "시정", "국가배상", "권리구제", "보충성"],
    55: ["공무원", "직무상 불법행위", "손해", "국민", "법률", "국가", "공공단체"],
    56: ["국가배상법", "배상청구", "가해공무원", "직접", "배상청구", "헌법"],
    57: ["공무원", "위법행위", "방지", "피해자", "권리구제", "직접청구"],
    58: [
        "공무원",
        "경과실",
        "손해",
        "국가",
        "기관",
        "고의중과실",
        "직접",
        "손해배상책임",
    ],
    60: ["기판력", "소송물", "법원", "판단", "확정", "구속력"],
    61: ["취소소송", "소송물", "처분", "위법성", "취소판결", "기판력"],
    62: ["전부기판력 긍정설", "협의의 행위위법성설", "기판력", "선행판결"],
    63: ["제한적 긍정설", "광의의 행위위법성설", "인용판결", "기각판결"],
    64: ["항고소송", "취소", "기판력", "행정처분", "공무원", "고의과실", "불법행위"],
    65: ["민사법원", "독자적", "법령위반", "국가배상청구요건", "심리"],
    68: ["영조물", "강학", "공물", "공공", "목적", "유체물", "물적설비"],
    70: ["설치", "관리", "하자", "공물", "용도", "안전성", "방호조치의무"],
    71: ["영조물", "책임", "공무원", "귀책사유", "무과실 책임"],
    73: ["행정주체", "공물", "공중", "사용", "안전조치", "법적의무", "무과실책임"],
    74: ["영조물", "설치관리", "하자", "안전성", "사회통념", "방호조치의무"],
    75: ["영조물", "물적설비", "물리적", "외형적", "흠결", "공해", "소음"],
    76: ["영조물", "공공", "목적", "이용", "수인할 것", "피해", "기능적 하자"],
    77: ["수인한도", "침해", "권리", "공공성", "환경", "종합적", "판단"],
    78: ["공물", "설치", "관리", "하자", "공무원", "직무상 불법행위", "배상청구"],
    80: ["국가", "지자체", "배상주체", "비용부담자", "배상책임"],
    81: ["지자체", "국가", "기관위임사무", "비용", "부담주체", "배상책임"],
    82: ["사무", "관리주체", "사무귀속자", "비용", "부담", "손해배상비용"],
    85: ["경찰공무원", "경찰업무", "위험성", "경찰조직", "전투경찰순경"],
    86: ["전투", "훈련", "직무집행", "일반 직무집행", "배상책임", "제한"],
    87: ["군인", "민간인", "공동불법행위", "피해", "배상", "구상"],
    88: ["헌법재판소", "국가배상법", "구상권", "일반국민", "국가", "차별", "재산권"],
    89: [
        "대법원",
        "한정위헌결정",
        "헌법재판소",
        "민간인",
        "피해자",
        "손해배상",
        "구상",
    ],
    96: ["공공필요", "재산권", "수용", "사용", "제한", "보상", "법률"],
    98: ["손실보상", "적법", "공행정작용", "재산권", "침해", "특별희생", "행정주체"],
    99: ["판례", "토지보상법", "손실보상금", "당사자소송", "공권"],
    100: ["공공필요", "공익목적", "재산권", "제한", "공익", "이익형량"],
    103: ["손실보상", "규정", "헌법", "근거", "손실보상청구"],
    104: ["헌법", "직접효력설", "위헌무효설", "보상입법부작위위헌설"],
    105: ["판례", "규정", "유추적용", "손실보상청구", "헌법재판소", "개발제한구역"],
    107: ["특별희생", "규정", "유추적용", "손실보상청구"],
    108: ["경계이론", "재산권", "수용", "사용", "제한", "특별희생", "보상"],
    109: ["분리이론", "재산권", "수용", "침해규정", "손실보상규정", "사회적 기속"],
    110: [
        "공행정작용",
        "침해",
        "특별희생",
        "사회적 기속",
        "재산권",
        "보호가치",
        "수인한도",
    ],
    111: ["목적위배설", "사적효용설", "수인한도설", "종합적", "판단"],
}


def create_quiz_final(source_path, output_dir):
    """
    원본 파일을 복사한 후 문단별로 빈칸 변경
    """
    filename = os.path.splitext(os.path.basename(source_path))[0]

    # 1. 원본 파일 복사
    question_path = os.path.join(output_dir, f"{filename}_문제.docx")
    shutil.copy2(source_path, question_path)

    # 2. 복사된 파일 열기
    doc = docx.Document(question_path)

    # 3. 각 문단에서 빈칸 변경
    for para_idx, blank_words in blank_words_by_paragraph.items():
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            for run in para.runs:
                for word in blank_words:
                    if word in run.text:
                        run.text = run.text.replace(word, "________")

    # 4. 모든 표에서도 빈칸 변경
    all_blank_words = []
    for words in blank_words_by_paragraph.values():
        all_blank_words.extend(words)
    all_blank_words = list(set(all_blank_words))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        for word in all_blank_words:
                            if word in run.text:
                                run.text = run.text.replace(word, "________")

    # 5. 저장
    doc.save(question_path)

    # 6. 답안 파일 생성
    answer_path = os.path.join(output_dir, f"{filename}_답안.docx")
    answer_doc = docx.Document()

    # 답안 제목 (인코딩 문제 해결을 위해 기본 폰트 사용)
    title_para = answer_doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{filename} 답안")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Arial"

    # 원본 파일 열어서 문단 제목 추출
    source_doc = docx.Document(source_path)

    # 제목 매핑 (인덱스 -> 제목)
    title_map = {}
    current_title = ""
    for i, para in enumerate(source_doc.paragraphs):
        if para.style.name.startswith("Heading"):
            current_title = para.text.strip()
        title_map[i] = current_title

    # 답안 작성
    answer_num = 1
    for para_idx, blank_words in sorted(blank_words_by_paragraph.items()):
        if para_idx < len(source_doc.paragraphs):
            # 문단 제목 가져오기
            title = title_map.get(para_idx, f"문단 {para_idx}")

            # 제목 출력
            title_para = answer_doc.add_paragraph()
            title_run = title_para.add_run(f"■ {title}")
            title_run.bold = True
            title_run.font.size = Pt(12)
            title_run.font.name = "Arial"

            # 빈칸 단어 목록
            for word in blank_words:
                answer_para = answer_doc.add_paragraph()
                answer_run = answer_para.add_run(f"  {answer_num}. {word}")
                answer_run.font.name = "Arial"
                answer_num += 1

            # 빈 줄
            answer_doc.add_paragraph("")

    answer_doc.save(answer_path)

    print(
        f"생성 완료: {os.path.basename(question_path)}, {os.path.basename(answer_path)}"
    )


if __name__ == "__main__":
    base_dir = r"D:\code proj\word proj"

    # 국가배상, 손실보상만 테스트
    source = os.path.join(base_dir, "사례_국가배상, 손실보상.docx")
    create_quiz_final(source, base_dir)
